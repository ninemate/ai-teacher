import hashlib
import re
from collections.abc import Iterable
from dataclasses import dataclass
from functools import lru_cache
from pathlib import Path

import fitz
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from ebooklib import ITEM_DOCUMENT, epub
from lingua import Language, LanguageDetectorBuilder

SUPPORTED_EXTENSIONS = {".pdf", ".epub", ".txt", ".md", ".docx"}
UNSUPPORTED_BUT_KNOWN = {".mobi", ".azw", ".azw3"}


@dataclass
class TextSegment:
    text: str
    locator: str


@dataclass
class ParsedDocument:
    title: str | None
    author: str | None
    language: str | None
    segments: list[TextSegment]


def hash_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def normalize_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


_LINGUA_LANGUAGES = (
    Language.HUNGARIAN, Language.ENGLISH, Language.GERMAN,
    Language.FRENCH, Language.ITALIAN, Language.ROMANIAN,
    Language.SLOVAK, Language.CZECH, Language.POLISH,
    Language.RUSSIAN, Language.CROATIAN, Language.SERBIAN,
    Language.TURKISH, Language.SPANISH, Language.DUTCH,
    Language.SWEDISH, Language.BOKMAL,
    Language.DANISH, Language.FINNISH,
)


@lru_cache(maxsize=1)
def _detector():
    return LanguageDetectorBuilder.from_languages(*_LINGUA_LANGUAGES).with_preloaded_language_models().build()


def detect_language(text: str) -> str | None:
    sample = text[:2000].strip()
    if len(sample) < 50:
        return None
    try:
        lang = _detector().detect_language_of(sample)
        return lang.iso_code_639_1.name.lower() if lang else None
    except Exception:  # noqa: BLE001 - best-effort detection, must never break ingestion
        return None


def parse_document(path: Path) -> ParsedDocument:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return parse_pdf(path)
    if suffix == ".epub":
        return parse_epub(path)
    if suffix in {".txt", ".md"}:
        return parse_plaintext(path)
    if suffix == ".docx":
        return parse_docx(path)
    if suffix in UNSUPPORTED_BUT_KNOWN:
        raise ValueError(
            f"{suffix} is not enabled in the MVP parser. Convert to EPUB/PDF/TXT or add Calibre-based conversion later."
        )
    raise ValueError(f"Unsupported file type: {suffix}")


def _ocr_page(pixmap: fitz.Pixmap, lang: str) -> str:
    import pytesseract
    from PIL import Image

    if pixmap.n != 3:
        pix = fitz.Pixmap(fitz.csRGB, pixmap)
    else:
        pix = pixmap
    img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
    result = pytesseract.image_to_string(img, lang=lang)
    if pix is not pixmap:
        del pix
    return normalize_text(result)


def parse_pdf(path: Path) -> ParsedDocument:
    from teacher_common.config import get_settings
    cfg = get_settings()
    segments: list[TextSegment] = []
    with fitz.open(path) as pdf:
        metadata = pdf.metadata or {}
        for index, page in enumerate(pdf, start=1):
            text = normalize_text(page.get_text("text"))
            is_ocr = False
            if (not text or len(text.strip()) < 50) and cfg.ocr_enabled:
                pix = page.get_pixmap(dpi=cfg.ocr_dpi)
                ocr_text = _ocr_page(pix, cfg.ocr_language)
                if ocr_text:
                    is_ocr = True
                    text = ocr_text
            if text:
                locator = f"oldal {index} (OCR)" if is_ocr else f"oldal {index}"
                segments.append(TextSegment(text=text, locator=locator))
    combined = "\n\n".join(segment.text for segment in segments)
    return ParsedDocument(
        title=metadata.get("title") or path.stem,
        author=metadata.get("author"),
        language=detect_language(combined),
        segments=segments,
    )


def parse_epub(path: Path) -> ParsedDocument:
    book = epub.read_epub(str(path))
    segments: list[TextSegment] = []
    title = _first_metadata(book, "DC", "title") or path.stem
    author = _first_metadata(book, "DC", "creator")
    for item in book.get_items_of_type(ITEM_DOCUMENT):
        soup = BeautifulSoup(item.get_body_content(), "html.parser")
        text = normalize_text(soup.get_text("\n"))
        if text:
            label = item.get_name() or "szakasz"
            segments.append(TextSegment(text=text, locator=f"fejezet {label}"))
    combined = "\n\n".join(segment.text for segment in segments)
    return ParsedDocument(
        title=title,
        author=author,
        language=detect_language(combined),
        segments=segments,
    )


def parse_plaintext(path: Path) -> ParsedDocument:
    text = normalize_text(path.read_text(encoding="utf-8", errors="ignore"))
    return ParsedDocument(
        title=path.stem,
        author=None,
        language=detect_language(text),
        segments=[TextSegment(text=text, locator="szövegtörzs")] if text else [],
    )


def parse_docx(path: Path) -> ParsedDocument:
    doc = DocxDocument(str(path))
    parts = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    text = normalize_text("\n\n".join(parts))
    return ParsedDocument(
        title=path.stem,
        author=None,
        language=detect_language(text),
        segments=[TextSegment(text=text, locator="dokumentum")] if text else [],
    )


def chunk_segments(segments: Iterable[TextSegment], chunk_size: int, chunk_overlap: int) -> list[TextSegment]:
    chunks: list[TextSegment] = []
    for segment in segments:
        text = normalize_text(segment.text)
        if not text:
            continue
        start = 0
        while start < len(text):
            end = min(len(text), start + chunk_size)
            if end < len(text):
                split = text.rfind(" ", start, end)
                if split > start + 100:
                    end = split
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(TextSegment(text=chunk, locator=segment.locator))
            if end >= len(text):
                break
            start = max(0, end - chunk_overlap)
    return chunks


def iter_library_files(root: Path, extensions: Iterable[str]) -> Iterable[Path]:
    allowed = {item.lower() for item in extensions}
    for path in sorted(root.rglob("*")):
        if path.is_file() and path.suffix.lower() in allowed.union(UNSUPPORTED_BUT_KNOWN):
            yield path


def _first_metadata(book, namespace: str, key: str) -> str | None:
    values = book.get_metadata(namespace, key)
    if not values:
        return None
    return values[0][0]
